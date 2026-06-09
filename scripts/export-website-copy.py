#!/usr/bin/env python3
"""Export all Sovereign Group website copy to a Word document."""

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from pathlib import Path

OUTPUT = Path(__file__).resolve().parents[1] / "Sovereign Group Website Copy.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.all_caps = True


def add_body(doc, text):
    doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_stat_row(doc, stats):
    for stat in stats:
        doc.add_paragraph(f"{stat['value']} — {stat['label']}")


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_heading("Sovereign Group Website Copy", 0)
    title.alignment = 0
    doc.add_paragraph(
        "Complete copy export for Retirement 247 and Sovereign Capital. "
        "Generated from packages/content and site components."
    )
    doc.add_paragraph()

    # ─── RETIREMENT 247 ───────────────────────────────────────────────
    add_heading(doc, "Retirement 247", 1)

    add_heading(doc, "Site metadata", 2)
    add_body(doc, "Default title: Retirement 247 · Life Rights Expertise. Always On.")
    add_body(doc, "Meta description: South Africa's specialist life rights advisory and sales platform for developers, landowners, funds, and REITs.")
    add_body(doc, "Open Graph title: Retirement 247")
    add_body(doc, "Open Graph description: Life rights expertise. Always on.")
    add_body(doc, "Site URL: https://retirement247.co.za")

    add_heading(doc, "Navigation", 2)
    add_bullets(doc, ["Why Life Rights", "Services", "Who We Serve", "Team", "Speak to Us"])

    add_heading(doc, "Footer", 2)
    add_body(doc, "Wordmark: retirement 247")
    add_body(doc, "Tagline: Strategy · Design · Marketing · Sales")
    add_body(doc, "Group line: A Sovereign Capital Group company")
    add_body(doc, "Link: sovereigncapital.co.za")
    add_body(doc, "Copyright: © Retirement 24Seven")

    add_heading(doc, "Homepage", 2)
    add_label(doc, "Hero eyebrow")
    add_body(doc, "Life Rights Advisory · South Africa & Mauritius")
    add_label(doc, "Headline")
    add_body(doc, "Life rights expertise. Always on.")
    add_label(doc, "Headline emphasis")
    add_body(doc, "Always on.")
    add_label(doc, "Intro")
    add_body(doc, "We work with developers, landowners, funds, and REITs entering senior living, without the cost of learning the model the hard way.")
    add_label(doc, "Mission quote")
    add_body(doc, "An exceptional lived experience for life right holders, and a well-managed asset for the investor.")
    add_label(doc, "Primary CTA")
    add_body(doc, "Speak to Us → /contact")
    add_label(doc, "Secondary CTA")
    add_body(doc, "Explore our services → #services")

    add_label(doc, "Market statistics")
    add_stat_row(doc, [
        {"value": "5.4M", "label": "South Africans over 60 today"},
        {"value": "2%", "label": "In formal retirement housing"},
        {"value": "5–10yr", "label": "Waiting lists at top villages"},
        {"value": "R80–100B", "label": "Total market value"},
    ])

    add_label(doc, "Why Life Rights (homepage section)")
    add_body(doc, "Headline: Why the life rights model is structurally different")
    add_body(doc, "Pull quote: The life right model, when ethically applied, offers a fair exchange, and a compounding long-term asset for the owner.")
    add_bullets(doc, [
        "A life right grants the purchaser the right to occupy a unit for life, protected by the Housing Development Schemes for Retired Persons Act No. 65 of 1988. The developer retains ownership of the underlying asset in perpetuity.",
        "This creates a fundamentally different commercial structure from sectional title or freehold development: an interest-free loan at entry, profit on every resale, and 8–10% annual unit churn from Year 5 creating a recurring, inflation-linked income stream.",
    ])
    add_body(doc, "Returns panel label: Developer returns at a glance")
    add_stat_row(doc, [
        {"value": "~20%", "label": "Return on cost, first-round sales (greenfield)"},
        {"value": "8–10%", "label": "Annual unit churn from Year 5 (recurring income)"},
        {"value": "100%", "label": "Resale profit retained by asset owner"},
        {"value": "3%", "label": "Effective tax rate on life right income (Brummeria)"},
    ])
    add_body(doc, "CTA: Read the full model → /why-life-rights")

    add_label(doc, "Editorial image caption")
    add_body(doc, "Community built for the long horizon.")

    add_label(doc, "Commercial pillars")
    pillars = [
        ("Asset Management", "Levy, finance, governance, and compliance: the asset stays institutional-grade.", ["Levy collections", "Financial management", "Governance", "Compliance"]),
        ("Operations Management", "Estate management, care, and daily life: the lived experience residents feel.", ["Estate management", "Care services", "Food & beverage", "Social activities"]),
        ("Strategic Advisory", "Acquisition, design, sales, and exit: the decisions that determine whether the numbers work.", ["Site acquisition", "Village design", "Sales management", "Strategic exit"]),
    ]
    for title, outcome, chips in pillars:
        doc.add_paragraph(f"{title}: {outcome}", style="List Bullet")
        doc.add_paragraph(f"Keywords: {', '.join(chips)}")

    add_label(doc, "Services section")
    add_body(doc, "Headline: End-to-end support. From concept to exit.")
    add_body(doc, "Intro: We are not a marketing agency. We are a specialist operating partner, embedded in your project from the first strategic conversation through to a fully functioning, self-sustaining retirement village. We can engage at any stage, on any mandate, at any scale.")

    services = [
        ("01", "Value Proposition Development", "The strategic brief that guides every downstream decision: product type, pricing, care model, and competitive positioning.", "We define the strategic brief that guides every downstream decision: product type, pricing architecture, care model, target market, and competitive positioning. The document that makes every subsequent conversation shorter and sharper."),
        ("02", "Village Design Advisory", "Commercially grounded input on site plans, unit mix, amenity programming, and care facility integration.", "We provide experienced, commercially grounded input to your architects on site development plans, unit mix, unit sizing, shared amenity programming, care facility integration, and the physical conditions that create a genuine sense of community."),
        ("03", "Financial Modelling", "Life rights financial models grounded in real transaction experience: pricing, churn, regrant revenue, and sensitivity scenarios.", "We build the life rights financial model, covering pricing strategy, churn forecasts, regrant revenue projections, levy modelling, sensitivity scenarios, and operating budgets, providing reliable assumptions grounded in real transaction experience, not theory."),
        ("04", "Life Right Agreement Advisory", "Operational feedback on the Life Right Agreement, protecting resident rights and owner interests under the Act.", "We provide experienced feedback to your legal team on the Life Right Agreement, flagging operational pitfalls that only become visible after residents move in, and ensuring the document protects both resident rights and owner interests under the Act."),
        ("05", "Governance Structure", "Governance models that satisfy resident participation without adversarial HOA dynamics.", "We advise on governance models that satisfy residents' legitimate desire for participation without creating the adversarial dynamics of HOAs and body corporates. Getting this right at the outset is one of the most consequential decisions in any life rights project."),
        ("06", "Sales & Marketing", "Launch strategy, sales team selection, and active sales management for senior living, a fundamentally different discipline.", "We design and manage the sales process, from launch strategy and sales team selection through to active sales management, coaching, and reporting. Senior living sales requires a fundamentally different approach from standard property sales; we train your team accordingly."),
        ("07", "Staffing & Operations", "Operating model design, organisational structure, and hands-on coaching through the first months of operation.", "We help you determine the optimal operating model, whether outsourced, internal, or hybrid, and guide the organisational structure and selection of senior staff. We provide hands-on coaching through the first months of operation to ensure a smooth transition."),
        ("08", "Care Services Oversight", "Person-centred care model design and care team capability building across the healthcare continuum.", "We believe in person-centred care that meets each resident's specific and evolving needs. We advise on care model design, oversee care team capability building, and ensure your village delivers the healthcare continuum residents expect and deserve."),
        ("09", "Strategic Exit", "Portfolio acquisition facilitation, in full or in part, by pension funds and institutional investors.", "When the time comes to exit the life rights business, we facilitate the acquisition of the portfolio, in full or in part, by a pension fund or institutional investor. We know the buyers, we understand the valuation mechanics, and we manage the transition."),
    ]
    for num, title, summary, body in services:
        add_heading(doc, f"{num} {title}", 3)
        add_body(doc, f"Summary: {summary}")
        add_body(doc, body)

    add_label(doc, "Who We Serve")
    add_body(doc, "Headline: Built for principals with a long-term orientation")
    add_body(doc, "Intro: We take mandates we can stand behind for 30 years: landowners who care about their legacy, developers who understand that a retirement village is a long-term commitment, and investors who want institutional-grade returns without operational shortcuts.")
    audiences = [
        ("Landowners & Families", "I have the land and the legacy. I need a partner who won't treat this as a short-cycle development.", "Generational landowners, from wine estates and farm holdings to coastal properties, who want to unlock long-term value without a short-cycle developer's exit strategy. We structure the partnership and manage the complexity.", "See landowner benefits"),
        ("Developers & Investors", "The numbers work on paper. I need someone who's done this before the residents move in.", "Developers entering retirement for the first time, and investors building senior living portfolios who need specialist advisory alongside their capital. We provide the operational intelligence that makes the numbers real.", "See developer benefits"),
        ("Funds & REITs", "We own the asset. We don't want to run the village, but we need it run properly.", "Institutional funds and REITs who own retirement assets but lack the appetite for day-to-day asset management and resident liaison. We take on the operational complexity so they can focus on portfolio performance.", "See fund benefits"),
    ]
    for title, quote, body, cta in audiences:
        add_heading(doc, title, 3)
        add_body(doc, f"Quote: \"{quote}\"")
        add_body(doc, body)
        add_body(doc, f"CTA: {cta}")

    add_label(doc, "Process steps")
    add_body(doc, "Headline: How we engage")
    add_bullets(doc, [
        "01 Assess — Site, market, and mandate fit, before any commitment.",
        "02 Structure — Model, agreements, governance, and financials, built on real transaction experience.",
        "03 Operate — Launch, sales, resident experience, and ongoing asset management.",
    ])

    add_label(doc, "Closing CTA")
    add_body(doc, "Headline: The life rights model suits your site. Let's run the numbers.")
    add_body(doc, "Subline: Conversations begin with a site, a mandate, and a long-term orientation.")
    add_body(doc, "Primary CTA: Speak to Us → /contact")
    add_body(doc, "Secondary CTA: Meet the team → /team")

    page_break(doc)

    add_heading(doc, "Why Life Rights page", 2)
    add_body(doc, "Page title: Why Life Rights")
    add_body(doc, "Meta description: Why the life rights model is structurally different, for asset owners, developers, and funds.")
    add_body(doc, "Eyebrow: Why Life Rights · Retirement 247")
    add_body(doc, "Hero headline: Why the life rights model is structurally different")
    add_body(doc, "Intro: The structural case for life rights as a long-duration asset class, for owners, developers, and funds.")
    add_body(doc, "Opportunity statement: Active mandates across the Winelands, the Garden Route, and into Mauritius, secured through direct relationship, not competitive pitch.")
    add_bullets(doc, [
        "A life right grants the purchaser the right to occupy a unit for life, protected by the Housing Development Schemes for Retired Persons Act No. 65 of 1988. The developer retains ownership of the underlying asset in perpetuity.",
        "This creates a fundamentally different commercial structure from sectional title or freehold development. The developer receives an interest-free loan at entry, sufficient to settle land and development debt, and retains the profit on every resale as the village matures. At maturity, life right unit churn of 8–10% per annum creates a recurring, inflation-linked income stream that grows in real terms with every cycle.",
        "No transfer duty. No VAT. No bad debt. No special levies. A tax regime treated favourably under the Brummeria judgment. And a resident base that is contractually aligned with the asset owner's long-term interests.",
    ])
    add_body(doc, "Image caption: Care with authority. Security with warmth.")
    add_body(doc, "Model section headline: For the asset owner and the resident")
    add_body(doc, "Accordion — For the Asset Owner:")
    add_bullets(doc, [
        "Developer return on cost of ~20% on first-round greenfield sales",
        "Interest-free loan from the purchaser at registration, sufficient to settle land and development debt",
        "Once-off Brummeria tax treatment on entry; ongoing income taxed at approximately 3%",
        "Repayment of the loan only on resale and registration of a replacement life right",
        "Profit on every resale retained by the asset owner as a realised fair value gain",
        "Life right assets held on balance sheet, appreciating annually (unrealised fair value gain)",
        "8–10% unit churn per annum from Year 5, recurring, compounding annuity income",
        "Operating costs recovered via resident levies once fully occupied",
        "No bad debt: the owner may access life right capital to settle levy arrears",
    ])
    add_body(doc, "Accordion — For the Resident:")
    add_bullets(doc, [
        "Security of tenure for life under the Housing Development Schemes for Retired Persons Act No. 65 of 1988",
        "No transfer duty, no VAT, no bond registration fees, a significant saving vs sectional title",
        "No special levies, a statutory protection with real financial consequence",
        "Interests aligned with the asset owner, with no perverse incentives",
        "Managed retirement lifestyle with communal facilities and professional estate management",
        "Access to continuous healthcare solutions via on-site or integrated facilities",
        "Community and security, particularly meaningful when family is absent or distant",
        "80–100% of original purchase price returned to estate on termination",
        "75% resident consent required for any sale of the underlying property",
    ])

    add_heading(doc, "Services page", 2)
    add_body(doc, "Page title: Services")
    add_body(doc, "Meta description: End-to-end life rights support, from value proposition development through to strategic exit.")
    add_body(doc, "Eyebrow: Services · Retirement 247")
    add_body(doc, "Intro: Nine disciplines of specialist support, from first strategic conversation through to exit.")
    add_body(doc, "Alternate services headline: Every discipline. One operating partner.")

    add_heading(doc, "Who We Serve page", 2)
    add_body(doc, "Page title: Who We Serve")
    add_body(doc, "Meta description: Landowners, developers, funds, and REITs entering senior living.")
    add_body(doc, "Eyebrow: Who We Serve · Retirement 247")

    add_heading(doc, "Team page", 2)
    add_body(doc, "Page title: Team")
    add_body(doc, "Meta description: Decades of life rights experience. One focused mandate.")
    add_body(doc, "Eyebrow: Team · Retirement 247")
    add_body(doc, "Headline: Decades of experience. One focused mandate.")
    add_body(doc, "Intro: Principals with decades of life rights development, operations, and transaction experience.")
    add_body(doc, "Secondary CTA: Explore our services → /services")
    for name, title, bio in [
        ("Gavin Vickers", "Director · Retirement 24Seven", "Specialist in development finance, deal structuring, and life rights capital. Former corporate banking roles at RMB Private Banking, Nedbank Commercial Asset Based Finance, and Standard Bank. Co-developer of Mt Prospect retirement estate in Constantia. Founding principal of RezProp (sold to National Real Estate, 2023). Co-founder and Director of Retirement 24Seven."),
        ("Arthur Case", "CEO · Sovereign Senior Living", "Former CEO of Evergreen Lifestyles (Amdec Group). Developed, commissioned and operated 7 retirement villages over 12 years. Career spanning Anglo American, Nedbank, Southern Life Association, and Deloitte across South Africa and the Middle East. Former Managing Director of MediCor Hospital Group (sold to MediClinic). The most operationally experienced life rights executive in South Africa."),
    ]:
        add_heading(doc, name, 3)
        add_body(doc, title)
        add_body(doc, bio)

    add_heading(doc, "Contact page", 2)
    add_body(doc, "Page title: Contact")
    add_body(doc, "Meta description: Speak to Retirement 247 about life rights advisory and sales.")
    add_body(doc, "Eyebrow: Contact · Retirement 247")
    add_body(doc, "Headline: The life rights model suits your site. Let's run the numbers.")
    add_body(doc, "Intro: Whether you have a site under consideration or an asset that needs better management, start with a conversation.")
    add_body(doc, "Pull quote: The life rights model suits this site well. We have run the numbers and can walk you through three structuring options at your convenience.")
    add_body(doc, "Advisory Enquiries: info@retirement24seven.co.za")
    add_body(doc, "Office: Somerset West, Western Cape, South Africa")
    add_body(doc, "Director: Gavin Vickers")
    add_body(doc, "CTA title: Start with a conversation")
    add_body(doc, "CTA body: Whether you have a site under consideration, a development in feasibility, or an existing retirement asset that needs better management, we are ready to engage. We will tell you honestly whether the life rights model is the right structure, and what it would take to make it work.")
    add_body(doc, "Consortium block label: Part of the Sovereign Capital Group")
    add_body(doc, "Consortium headline: Advisory is only part of what we bring.")
    add_body(doc, "Consortium body: Retirement 24Seven operates as the specialist advisory and sales arm of the Sovereign Capital Group, a private capital consortium that also provides insurance-backed operating infrastructure through Absolute Life, and access to institutional finance through our collaboration with Fedgroup. For the right mandate, we are not just advisors. We are co-investors.")
    add_bullets(doc, [
        "Sovereign Capital — Holding Co · Finance",
        "Absolute Life — Operating Platform",
        "Lombard Insurance — Risk Partner",
        "Fedgroup — Financial Collaborator",
        "Retirement 24Seven — Advisory · Sales",
    ])

    add_heading(doc, "Image alt text & captions", 2)
    images = [
        ("Hero water", "Calm navy water at golden hour", None),
        ("Editorial leaves", "Soft green leaves in natural light", "Community built for the long horizon."),
        ("Editorial team", "A collaborative team of professionals planning together", None),
        ("Care hands", "An elderly person's hand gently held, a quiet gesture of care and reassurance", "Care with authority. Security with warmth."),
        ("Community joy", "Older adults smiling together outdoors, enjoying community and connection", "Community built for the long horizon."),
        ("Companions walking", "An elderly couple walking together through a sunlit garden", None),
        ("Friends laughing", "Two older friends laughing and enjoying each other's company", None),
    ]
    for name, alt, caption in images:
        add_body(doc, f"{name}: {alt}")
        if caption:
            add_body(doc, f"Caption: {caption}")

    add_heading(doc, "404 page", 2)
    add_body(doc, "Headline: Page not found")

    page_break(doc)

    # ─── SOVEREIGN CAPITAL ────────────────────────────────────────────
    add_heading(doc, "Sovereign Capital", 1)

    add_heading(doc, "Site metadata", 2)
    add_body(doc, "Default title: Sovereign Capital — Capital with Conscience")
    add_body(doc, "Meta description: A private capital group deploying patient, long-duration capital into South Africa's life rights retirement sector.")
    add_body(doc, "Open Graph description: Capital with conscience. Care with authority.")
    add_body(doc, "Site URL: https://sovereigncapital.co.za")

    add_heading(doc, "Navigation", 2)
    add_bullets(doc, ["About", "Developments", "Senior Living", "Money", "Private Equity", "Impact", "Leadership", "Contact"])

    add_heading(doc, "Footer", 2)
    add_body(doc, "Wordmark: SOVEREIGN CAPITAL")
    add_body(doc, "Locations: Somerset West · Stellenbosch · Mauritius")
    add_body(doc, "Links: About, Contact, Retirement 247")
    add_body(doc, "Copyright: © Sovereign Capital (Pty) Ltd")

    add_heading(doc, "Homepage", 2)
    add_body(doc, "Tagline: Capital with conscience. / Care with authority.")
    add_body(doc, "Locations: Somerset West · Stellenbosch · Mauritius")
    add_stat_row(doc, [
        {"value": "R9B+", "label": "Life Rights Sales"},
        {"value": "7+", "label": "Active Projects"},
        {"value": "R30B", "label": "Financial Partner AUM"},
        {"value": "2", "label": "Markets: SA & Mauritius"},
    ])
    add_label(doc, "Who We Are")
    add_body(doc, "Label: WHO WE ARE")
    add_body(doc, "Headline: A private capital group built for the long horizon")
    add_bullets(doc, [
        "Sovereign Capital is the institutional holding entity and master brand of a specialist group operating at the intersection of senior living development, structured finance, and life rights advisory. We are based in Somerset West, Western Cape, with active mandates extending to Mauritius.",
        "We structure, fund, and operate life rights retirement villages across the Western Cape and beyond — partnering with landowners, family offices, REITs, and developers who want institutional-grade execution without surrendering long-term control of their assets.",
    ])
    add_body(doc, "Pull quote: We are asset custodians and operators — not short-cycle developers.")
    add_body(doc, "Follow-up: The life rights model, when ethically applied, creates a compounding, perpetual income engine for the asset owner while delivering genuine security and community to residents. This alignment of interests is why we focus on it exclusively.")
    add_body(doc, "Market context: South Africa's over-60 population stands at 5.4 million today and is projected to reach 12 million by 2050. Only 2% of seniors currently reside in formal housing. The structural gap is vast, the timing is now, and the patient capital we deploy is ideally suited to capture it.")
    add_body(doc, "Business units headline: Five disciplines. One coordinated structure.")
    add_body(doc, "Contact pull quote: Conversations begin quietly.")

    add_heading(doc, "About page", 2)
    add_body(doc, "Page title: About")
    add_body(doc, "Meta description: Group identity, structure, and philosophy.")
    add_body(doc, "Headline: Group identity, structure, and philosophy")
    add_body(doc, "Geography label: GEOGRAPHY")
    add_body(doc, "Geography headline: Western Cape rooted. Indian Ocean reaching.")
    add_body(doc, "South Africa: Our primary market. Active projects span the Winelands, the Atlantic Seaboard, the Garden Route, and the Cape West Coast — concentrated in the Western Cape's semi-gration corridor where the structural case for premium life rights is strongest. Sovereign Capital's principal office is in Somerset West.")
    add_body(doc, "Mauritius: We hold a signed mandate with Ferney Limited (Luke Maurel) for Mauritius' first Life Right retirement village within the Ferney Smart City — a landmark development addressing a market with 282 retirement units serving a rapidly ageing population of 184,000 seniors. Senior Living Fund (Pty) Ltd leads this mandate.")

    add_heading(doc, "Business units", 2)
    units = [
        ("Sovereign Senior Living", "Life rights development and senior living operations at institutional scale.", "We identify, structure, and develop life rights retirement villages — from greenfield opportunity through to fully occupied, self-sustaining estate.", [
            "Sovereign Senior Living is the group's senior living development and operations platform. We manage the full value chain: land, design, feasibility, funding, sales, and operations — partnering with landowners, family offices, REITs, and developers who want institutional-grade execution without surrendering long-term control of their assets.",
            "Through Retirement 24Seven, our specialist advisory and sales arm, we bring decades of life rights expertise to every mandate — whether Sovereign-led or third-party.",
        ], [("01", "Life Rights Development", "Greenfield and brownfield life rights retirement villages across South Africa and Mauritius — from opportunity identification through to fully occupied estate."), ("02", "Senior Living Advisory", "Financial modelling, unit design, governance structures, sales strategy, and statutory compliance under the Housing Development Schemes for Retired Persons Act."), ("03", "Operations & Care", "Insurance-backed operating infrastructure through Absolute Life — person-centred care, estate management, and long-duration asset stewardship.")], "Retirement 24Seven → https://retirement247.co.za"),
        ("Sovereign Developments", "Institutional-grade life rights development across the Western Cape and beyond.", "We structure, fund, and develop life rights retirement villages — managing the full value chain from land through to operations.", [
            "Sovereign Developments leads the group's property development mandates — concentrated in the Western Cape's semi-gration corridor where the structural case for premium life rights is strongest.",
            "Our consortium model brings together capital, insurance-backed operating infrastructure, and specialist advisory under a single coordinated structure — institutional-grade governance with the agility of a specialist boutique.",
        ], [("01", "Project Structuring", "JV frameworks, MOUs, and HOA arrangements between landowners, developers, and capital providers — aligning interests across multi-phase, long-duration projects."), ("02", "Development Management", "End-to-end development oversight from feasibility through commissioning — design advisory, contractor management, and sales coordination."), ("03", "Institutional Partnerships", "Capital partner introductions, funder relationships, and co-investment structures for projects requiring institutional-grade financial partners.")], None),
        ("Sovereign Money", "Structured finance and bridging capital for long-duration property assets.", "Bridging and mezzanine finance for the broader development market — Tier 1 conventional bridging and Tier 2 development-linked facilities.", [
            "Sovereign Money operates a bridging and mezzanine finance unit providing short-term development capital to the broader property market. Backed by our institutional partners, we deploy Tier 1 conventional bridging and Tier 2 development-linked facilities where bank appetite is limited or too slow.",
            "Our finance capability extends the group's reach beyond life rights — serving developers who need patient, structurally sound capital with the governance discipline of an institutional consortium.",
        ], [("01", "Bridging Finance", "Short-term development capital — Tier 1 conventional bridging facilities backed by institutional partners including Fedgroup."), ("02", "Mezzanine Structures", "Tier 2 development-linked facilities where bank appetite is limited — structured to align with project milestones and long-duration asset profiles."), ("03", "Capital Structuring", "Deal structuring, capital stack optimisation, and institutional partner introductions for complex, multi-phase development mandates.")], None),
        ("Sovereign Private Equity", "Patient private equity capital for specialist property sectors and long-duration assets.", "Institutional co-investment and direct equity mandates across senior living, development, and adjacent property sectors.", [
            "Sovereign Private Equity deploys patient capital into specialist property sectors where structural demand, regulatory alignment, and operational depth create durable returns over generational horizons.",
            "We partner with family offices, institutional investors, and co-investment partners who share our conviction that disciplined governance and long-duration asset stewardship outperform short-cycle development models.",
        ], [("01", "Direct Equity Mandates", "Co-investment and direct equity participation in life rights developments, senior living platforms, and specialist property assets — structured for institutional governance and long-horizon returns."), ("02", "Portfolio Strategy", "Mandate design, asset allocation, and portfolio construction across the group's specialist sectors — aligning capital deployment with structural market opportunity."), ("03", "Institutional Partnerships", "Co-investment structures, LP/GP frameworks, and capital partner introductions for investors seeking exposure to South Africa's senior living and specialist property sectors.")], None),
        ("Sovereign Impact", "Purpose-driven development beyond life rights — student accommodation and community programmes.", "PBSA development at institutional scale and long-term community investment through the Ground Up Programme.", [
            "Sovereign Impact encompasses the group's purpose-driven initiatives — developments and programmes that extend our capital stewardship beyond senior living into adjacent sectors and community investment.",
            "These mandates reflect our conviction that patient capital and institutional discipline can serve both commercial returns and genuine social impact.",
        ], [("01", "Student Accommodation", "Through our PBSA division, we develop and structure purpose-built student accommodation at institutional scale — from premium Studenthood-branded models to NSFAS-aligned affordable housing — with feasibility analysis, investor decks, and funder introductions across South Africa's university precincts."), ("02", "Ground Up Programme", "The Ground Up Programme is Sovereign Senior Living's long-term commitment to young South African men aged 20–45. Built in partnership with C-Suite executive coach Justin King and mentor Tony Karam, it delivers emotional intelligence, financial literacy, and mentorship without preconditions and without end dates.")], None),
    ]
    for name, prop, desc, body_paras, caps, link in units:
        add_heading(doc, name, 3)
        add_body(doc, f"Proposition: {prop}")
        add_body(doc, f"Descriptor: {desc}")
        for para in body_paras:
            add_body(doc, para)
        for num, cap_title, cap_body in caps:
            add_body(doc, f"{num} {cap_title}: {cap_body}")
        if link:
            add_body(doc, f"External link: {link}")

    add_heading(doc, "Leadership page", 2)
    add_body(doc, "Page title: Leadership")
    add_body(doc, "Meta description: The principals behind the Sovereign Capital consortium.")
    add_body(doc, "Breadcrumb: Sovereign Capital → Leadership")
    add_body(doc, "Headline: The principals behind the consortium.")
    for name, title, bio in leadership():
        add_heading(doc, name, 3)
        add_body(doc, title)
        add_body(doc, bio)

    add_heading(doc, "Contact page", 2)
    add_body(doc, "Page title: Contact")
    add_body(doc, "Meta description: Conversations begin quietly — partner with Sovereign Capital.")
    add_body(doc, "Headline: Conversations begin quietly.")
    add_body(doc, "CTA title: Partner with us")
    add_body(doc, "Email: gavin@sovereigncapital.co.za")
    add_body(doc, "CTA body: We work with landowners with generational ambitions, developers seeking institutional co-investment, funds building senior living portfolios, and capital providers looking for structurally sound long-duration assets. If your horizon is long and your standards are high, the conversation is worth having.")
    add_body(doc, "Locations: Somerset West · Stellenbosch · Mauritius")

    add_heading(doc, "404 page", 2)
    add_body(doc, "Headline: Page not found")

    doc.save(OUTPUT)
    return OUTPUT


def consortium_intro():
    return "We do not assemble partners project by project. Our consortium is a standing, purpose-built structure with each entity contributing a distinct and essential capability. The result is institutional-grade governance with the agility of a specialist boutique."


def consortium_quote():
    return "Lombard's underwriting strength anchors the Sovereign Capital consortium. Fedgroup's capital capacity extends the reach of Retirement 24Seven into transactions that require institutional-grade financial partners alongside specialist advisory."


def consortium_entities():
    return [
        ("Sovereign Capital", "Holding Company · Consortium Lead", "Institutional governance, deal structuring, capital deployment, and bridging finance. The entity that partners encounter first and that holds the group together."),
        ("Absolute Life", "Senior Living JV · Sovereign Capital & Lombard", "Insurance-backed operating platform jointly owned by Sovereign Capital and Lombard Insurance. Purpose-built to develop, fund, and manage senior living assets at institutional scale."),
        ("Retirement 24Seven", "Life Rights Advisory · Sales", "Specialist life rights advisory and sales platform. Accepts mandates across all Sovereign Capital projects and from third-party developers, funds, and landowners."),
        ("Lombard Insurance", "Sovereign Capital Partner · Insurance & Risk", "Established South African insurer and co-owner of Absolute Life. Lombard brings embedded insurance underwriting, structured risk solutions, and the regulatory credibility essential for a life rights operating platform."),
        ("Fedgroup", "Financial Collaborator · Retirement 24Seven · R30B AUM", "One of South Africa's most respected independent financial services groups. Fedgroup collaborates with Retirement 24Seven on selected life rights development projects — providing institutional-grade funding capacity, long-term capital commitment, and the financial stability required for large-scale phased developments."),
    ]


def leadership():
    return [
        ("Lauren Peacock", "Executive Director · Sovereign Capital", "Executive Director of Sovereign Capital, responsible for group governance, strategic partnerships, and operational leadership across the Sovereign consortium. Lauren brings rigorous institutional discipline to a group operating across multiple complex, long-duration development mandates simultaneously."),
        ("Rhys Meredith", "Executive Director · Sovereign Capital", "Executive Director of Sovereign Capital, contributing senior leadership across the group's capital structuring, development partnerships, and institutional relationships. Rhys brings depth of experience to a consortium operating at the intersection of property development, structured finance, and long-duration asset management."),
        ("Arthur Case", "Chief Executive Officer · Sovereign Senior Living", "Former CEO of Evergreen Lifestyles (Amdec Group) — developed, commissioned, and operated 7 retirement villages over 12 years. Career spanning Anglo American, Nedbank, Southern Life Association, and Deloitte across South Africa and the Middle East. Former Managing Director of MediCor Hospital Group (sold to MediClinic). Leads Sovereign Senior Living, bringing unmatched operational depth to the group's life rights development platform."),
        ("Gavin Vickers", "Director · Retirement 24Seven", "Specialist in development finance, bridging and mezzanine structures, and strategic deal-making. Former corporate banking positions at RMB Private Banking, Nedbank Commercial Asset Based Finance, and Standard Bank. Co-developer of Mt Prospect retirement estate in Constantia. Founding principal of RezProp (sold to National Real Estate, 2023). Co-founder and Director of Retirement 24Seven."),
    ]


if __name__ == "__main__":
    path = build_document()
    print(f"Created: {path}")
